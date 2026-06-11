from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path


class DocumentExtractionError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


@dataclass
class ExtractedDocumentText:
    text: str
    file_type: str
    parser_version: str = "document_text_extractor_v0.1"
    page_count: int | None = None
    text_quality: str = "empty"
    warnings: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "text": self.text,
            "file_type": self.file_type,
            "parser_version": self.parser_version,
            "page_count": self.page_count,
            "text_quality": self.text_quality,
            "warnings": self.warnings,
        }


def extract_document_text(path: str | Path, filename: str | None = None) -> ExtractedDocumentText:
    source = Path(path)
    suffix = (Path(filename or source.name).suffix or source.suffix).lower()
    if suffix == ".txt":
        return _extract_txt(source)
    if suffix == ".docx":
        return _extract_docx(source)
    if suffix == ".pdf":
        return _extract_pdf(source)
    raise DocumentExtractionError("unsupported_legal_document_type", "仅支持 PDF、DOCX、TXT 合同或条款文件。")


def _quality(text: str, warnings: list[str]) -> str:
    length = len(text.strip())
    if length == 0:
        return "empty"
    if "needs_ocr" in warnings:
        return "needs_ocr"
    if length < 100:
        return "low"
    if length < 1000:
        return "medium"
    return "high"


def _extract_txt(path: Path) -> ExtractedDocumentText:
    raw = path.read_bytes()
    warnings: list[str] = []
    text = ""
    for encoding in ("utf-8", "gb18030"):
        try:
            text = raw.decode(encoding)
            if encoding != "utf-8":
                warnings.append("decoded_with_gb18030")
            break
        except UnicodeDecodeError:
            continue
    if not text:
        raise DocumentExtractionError("document_decode_failed", "TXT 文件编码无法识别，请另存为 UTF-8 后重试。")
    text = _normalize_text(text)
    if not text.strip():
        warnings.append("empty_text")
    return ExtractedDocumentText(text=text, file_type="txt", text_quality=_quality(text, warnings), warnings=warnings)


def _extract_docx(path: Path) -> ExtractedDocumentText:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError("missing_python_docx", "缺少 python-docx 依赖，请先安装 requirements.txt。") from exc

    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    text = _normalize_text("\n".join(chunks))
    warnings = ["empty_text"] if not text.strip() else []
    return ExtractedDocumentText(text=text, file_type="docx", text_quality=_quality(text, warnings), warnings=warnings)


def _extract_pdf(path: Path) -> ExtractedDocumentText:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError("missing_pypdf", "缺少 pypdf 依赖，请先安装 requirements.txt。") from exc

    reader = PdfReader(str(path))
    chunks: list[str] = []
    warnings: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
            warnings.append("page_extract_failed")
        if page_text.strip():
            chunks.append(page_text.strip())
    text = _normalize_text("\n".join(chunks))
    if not text.strip():
        warnings.append("needs_ocr")
    return ExtractedDocumentText(text=text, file_type="pdf", page_count=len(reader.pages), text_quality=_quality(text, warnings), warnings=warnings)


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact: list[str] = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank:
                compact.append("")
            previous_blank = True
            continue
        compact.append(line)
        previous_blank = False
    return "\n".join(compact).strip()
