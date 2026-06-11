from __future__ import annotations

import html
import re
import shutil
import tempfile
from dataclasses import dataclass, field
from pathlib import Path


class DocumentExtractionError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".html", ".htm", ".png", ".jpg", ".jpeg", ".webp"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}


@dataclass
class IngestedDocument:
    text: str
    file_type: str
    parser_version: str = "document_ingestion_v0.1"
    page_count: int | None = None
    text_quality: str = "empty"
    warnings: list[str] = field(default_factory=list)
    is_scanned_pdf: bool = False
    extraction_method: str = "unknown"
    pages_used: list[int] = field(default_factory=list)
    ocr_status: str = "not_needed"
    ocr_confidence: float | None = None
    attachments: list[dict] = field(default_factory=list)
    field_sources: dict = field(default_factory=dict)

    def as_dict(self) -> dict:
        return {
            "text": self.text,
            "file_type": self.file_type,
            "parser_version": self.parser_version,
            "page_count": self.page_count,
            "text_quality": self.text_quality,
            "warnings": self.warnings,
            "is_scanned_pdf": self.is_scanned_pdf,
            "extraction_method": self.extraction_method,
            "pages_used": self.pages_used,
            "ocr_status": self.ocr_status,
            "ocr_confidence": self.ocr_confidence,
            "attachments": self.attachments,
            "field_sources": self.field_sources,
        }


def parser_status() -> dict:
    python_deps = _ocr_python_deps()
    tesseract_path = shutil.which("tesseract")
    poppler_path = shutil.which("pdftoppm") or shutil.which("pdfinfo")
    if python_deps["available"] and tesseract_path and poppler_path:
        ocr_status = "available"
    elif not python_deps["available"]:
        ocr_status = "missing_python_deps"
    elif not tesseract_path:
        ocr_status = "missing_tesseract"
    else:
        ocr_status = "missing_poppler"
    next_actions = []
    if ocr_status == "missing_python_deps":
        next_actions.append("安装可选依赖：.venv/bin/pip install -r requirements-ocr.txt")
    if ocr_status == "missing_tesseract":
        next_actions.append("安装本地 tesseract，例如 macOS: brew install tesseract tesseract-lang")
    if ocr_status == "missing_poppler":
        next_actions.append("安装 PDF 渲染工具 poppler，例如 macOS: brew install poppler")
    if not next_actions:
        next_actions.append("可解析图片型 PDF、PNG、JPG 和 WEBP。")
    return {
        "ok": True,
        "parser": "document_ingestion_v0.1",
        "supported_formats": sorted(SUPPORTED_SUFFIXES),
        "default_max_pages": 5,
        "ocr": {
            "status": ocr_status,
            "python_deps": python_deps,
            "tesseract_present": bool(tesseract_path),
            "poppler_present": bool(poppler_path),
            "reserved_engines": ["paddleocr"],
        },
        "next_actions": next_actions,
    }


def ingest_document(path: str | Path, filename: str | None = None, max_pages: int = 5) -> IngestedDocument:
    source = Path(path)
    suffix = (Path(filename or source.name).suffix or source.suffix).lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DocumentExtractionError("unsupported_legal_document_type", "仅支持 PDF、DOCX、TXT、HTML 和常见图片文件。")
    if suffix == ".txt":
        return _extract_txt(source)
    if suffix == ".docx":
        return _extract_docx(source)
    if suffix in {".html", ".htm"}:
        return _extract_html(source)
    if suffix == ".pdf":
        return _extract_pdf(source, max_pages=max_pages)
    return _extract_image(source, suffix)


def ingest_bytes(raw: bytes, filename: str, max_pages: int = 5) -> IngestedDocument:
    suffix = Path(filename).suffix.lower() or ".bin"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        path = Path(tmp.name)
        tmp.write(raw)
    try:
        return ingest_document(path, filename, max_pages=max_pages)
    finally:
        path.unlink(missing_ok=True)


def _ocr_python_deps() -> dict:
    missing: list[str] = []
    for module in ["pdf2image", "pytesseract", "PIL"]:
        try:
            __import__(module)
        except ImportError:
            missing.append("Pillow" if module == "PIL" else module)
    return {"available": not missing, "missing": missing}


def _quality(text: str, warnings: list[str]) -> str:
    length = len(text.strip())
    if length == 0:
        return "empty"
    if "needs_ocr" in warnings or "ocr_low_confidence" in warnings:
        return "needs_ocr" if length < 100 else "low"
    if length < 100:
        return "low"
    if length < 1000:
        return "medium"
    return "high"


def _extract_txt(path: Path) -> IngestedDocument:
    raw = path.read_bytes()
    warnings: list[str] = []
    text = ""
    for encoding in ("utf-8", "gb18030"):
        try:
            text = raw.decode(encoding)
            if encoding != "utf-8":
                warnings.append("decoded_with_gb18030")
            break
        except UnicodeDecodeError:
            continue
    if not text:
        raise DocumentExtractionError("document_decode_failed", "TXT 文件编码无法识别，请另存为 UTF-8 后重试。")
    text = _normalize_text(text)
    if not text.strip():
        warnings.append("empty_text")
    return IngestedDocument(text=text, file_type="txt", text_quality=_quality(text, warnings), warnings=warnings, extraction_method="text", field_sources={"body": "txt"})


def _extract_docx(path: Path) -> IngestedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentExtractionError("missing_python_docx", "缺少 python-docx 依赖，请先安装 requirements.txt。") from exc

    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    text = _normalize_text("\n".join(chunks))
    warnings = ["empty_text"] if not text.strip() else []
    return IngestedDocument(text=text, file_type="docx", text_quality=_quality(text, warnings), warnings=warnings, extraction_method="docx_text", field_sources={"body": "docx"})


def _extract_html(path: Path) -> IngestedDocument:
    raw = path.read_bytes()
    text = ""
    for encoding in ("utf-8", "gb18030", "gbk"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if not text:
        text = raw.decode("utf-8", errors="ignore")
    body, attachments = html_to_text_and_attachments(text)
    warnings = ["empty_text"] if not body.strip() else []
    return IngestedDocument(
        text=body,
        file_type="html",
        text_quality=_quality(body, warnings),
        warnings=warnings,
        extraction_method="html_text",
        attachments=attachments,
        field_sources={"body": "html"},
    )


def _extract_pdf(path: Path, max_pages: int = 5) -> IngestedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError("missing_pypdf", "缺少 pypdf 依赖，请先安装 requirements.txt。") from exc

    reader = PdfReader(str(path))
    chunks: list[str] = []
    warnings: list[str] = []
    pages_used: list[int] = []
    limit = min(len(reader.pages), max_pages)
    for index, page in enumerate(reader.pages[:limit], start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
            warnings.append("page_extract_failed")
        if page_text.strip():
            chunks.append(f"[page {index}]\n{page_text.strip()}")
            pages_used.append(index)
    text = _normalize_text("\n".join(chunks))
    is_scanned = not text.strip()
    extraction_method = "pdf_text"
    ocr_status = "not_needed"
    ocr_confidence = None
    if is_scanned:
        warnings.append("needs_ocr")
        ocr_text, ocr_meta = _extract_pdf_ocr_optional(path, max_pages=limit)
        warnings.extend(ocr_meta["warnings"])
        ocr_status = ocr_meta["status"]
        ocr_confidence = ocr_meta["confidence"]
        pages_used = ocr_meta["pages_used"]
        extraction_method = "pdf_ocr" if ocr_text.strip() else "ocr_unavailable"
        if ocr_text.strip():
            text = _normalize_text(ocr_text)
            warnings.append("ocr_applied")
    elif len(reader.pages) > max_pages:
        warnings.append("page_limit_applied")
    return IngestedDocument(
        text=text,
        file_type="pdf",
        page_count=len(reader.pages),
        text_quality=_quality(text, warnings),
        warnings=_unique(warnings),
        is_scanned_pdf=is_scanned,
        extraction_method=extraction_method,
        pages_used=pages_used,
        ocr_status=ocr_status,
        ocr_confidence=ocr_confidence,
        field_sources={"body": extraction_method},
    )


def _extract_image(path: Path, suffix: str) -> IngestedDocument:
    ocr_text, ocr_meta = _extract_image_ocr_optional(path)
    warnings = ocr_meta["warnings"]
    if ocr_text.strip():
        warnings.append("ocr_applied")
    else:
        warnings.append("needs_ocr")
    text = _normalize_text(ocr_text)
    return IngestedDocument(
        text=text,
        file_type=suffix.lstrip("."),
        text_quality=_quality(text, warnings),
        warnings=_unique(warnings),
        is_scanned_pdf=False,
        extraction_method="image_ocr" if text else "ocr_unavailable",
        pages_used=[1] if text else [],
        ocr_status=ocr_meta["status"],
        ocr_confidence=ocr_meta["confidence"],
        field_sources={"body": "image_ocr" if text else "image"},
    )


def _extract_pdf_ocr_optional(path: Path, max_pages: int) -> tuple[str, dict]:
    status = parser_status()["ocr"]["status"]
    if status != "available":
        return "", {"status": status, "confidence": None, "pages_used": [], "warnings": ["ocr_unavailable", status]}
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        return "", {"status": "missing_python_deps", "confidence": None, "pages_used": [], "warnings": ["ocr_unavailable", "missing_python_deps"]}

    try:
        pages = convert_from_path(str(path), first_page=1, last_page=max_pages)
    except Exception:
        return "", {"status": "ocr_render_failed", "confidence": None, "pages_used": [], "warnings": ["ocr_render_failed"]}

    chunks: list[str] = []
    pages_used: list[int] = []
    warnings = ["ocr_attempted"]
    for index, page in enumerate(pages, start=1):
        try:
            page_text = pytesseract.image_to_string(page, lang="chi_sim+eng")
        except Exception:
            page_text = ""
            warnings.append("ocr_page_failed")
        if page_text.strip():
            chunks.append(f"[page {index}]\n{page_text.strip()}")
            pages_used.append(index)
    if not chunks:
        warnings.append("ocr_empty")
    text = "\n".join(chunks)
    return text, {"status": "success" if text.strip() else "ocr_empty", "confidence": _ocr_confidence(text), "pages_used": pages_used, "warnings": warnings}


def _extract_image_ocr_optional(path: Path) -> tuple[str, dict]:
    status = parser_status()["ocr"]["status"]
    if status != "available":
        return "", {"status": status, "confidence": None, "warnings": ["ocr_unavailable", status]}
    try:
        import pytesseract
    except ImportError:
        return "", {"status": "missing_python_deps", "confidence": None, "warnings": ["ocr_unavailable", "missing_python_deps"]}
    try:
        text = pytesseract.image_to_string(str(path), lang="chi_sim+eng")
    except Exception:
        return "", {"status": "ocr_image_failed", "confidence": None, "warnings": ["ocr_image_failed"]}
    warnings = ["ocr_attempted"]
    if not text.strip():
        warnings.append("ocr_empty")
    return text, {"status": "success" if text.strip() else "ocr_empty", "confidence": _ocr_confidence(text), "warnings": warnings}


def _ocr_confidence(text: str) -> float | None:
    cleaned = re.sub(r"\s+", "", text or "")
    if not cleaned:
        return None
    useful = sum(1 for char in cleaned if "\u4e00" <= char <= "\u9fff" or char.isalnum())
    return round(min(0.95, max(0.2, useful / max(1, len(cleaned)))), 2)


def html_to_text_and_attachments(raw: str) -> tuple[str, list[dict]]:
    attachments = []
    for match in re.finditer(r"<a[^>]+href=['\"]([^'\"]+)['\"][^>]*>(.*?)</a>", raw, re.I | re.S):
        label = _clean_text(match.group(2))[:80] or "附件"
        url = html.unescape(match.group(1))
        attachments.append({"label": label, "url": url, "file_type": Path(url).suffix.lower().lstrip(".") or "unknown"})
    without_scripts = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", raw, flags=re.I | re.S)
    text = re.sub(r"<[^>]+>", "\n", without_scripts)
    return _clean_text(html.unescape(text)), attachments


def _clean_text(value: str) -> str:
    value = re.sub(r"&nbsp;", " ", value)
    value = re.sub(r"\r", "\n", value)
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{2,}", "\n", value)
    return value.strip()


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact: list[str] = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank:
                compact.append("")
            previous_blank = True
            continue
        compact.append(line)
        previous_blank = False
    return "\n".join(compact).strip()


def _unique(values: list[str]) -> list[str]:
    result = []
    for value in values:
        if value and value not in result:
            result.append(value)
    return result
