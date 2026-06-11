from __future__ import annotations

import os
import tempfile
import unittest
from pathlib import Path

from backend.core.document_text_extractor import DocumentExtractionError, extract_document_text


class DocumentTextExtractorTests(unittest.TestCase):
    def test_txt_utf8_and_gb18030_extract(self):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            path = Path(tmp.name)
        try:
            path.write_bytes("借款合同\n约定由深圳市南山区人民法院管辖。".encode("gb18030"))
            result = extract_document_text(path, "合同.txt")
            self.assertIn("借款合同", result.text)
            self.assertEqual(result.file_type, "txt")
            self.assertIn("decoded_with_gb18030", result.warnings)
        finally:
            os.unlink(path)

    def test_docx_extracts_paragraphs_and_tables(self):
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = Path(tmp.name)
        try:
            doc = Document()
            doc.add_paragraph("个人借款合同")
            doc.add_paragraph("双方约定仲裁委员会仲裁。")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "送达地址"
            table.cell(0, 1).text = "合同载明地址"
            doc.save(str(path))
            result = extract_document_text(path, "合同.docx")
            self.assertIn("个人借款合同", result.text)
            self.assertIn("送达地址", result.text)
            self.assertEqual(result.file_type, "docx")
        finally:
            os.unlink(path)

    def test_blank_pdf_returns_needs_ocr_warning(self):
        from pypdf import PdfWriter

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            path = Path(tmp.name)
        try:
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            with path.open("wb") as handle:
                writer.write(handle)
            result = extract_document_text(path, "扫描合同.pdf")
            self.assertEqual(result.file_type, "pdf")
            self.assertEqual(result.text_quality, "empty")
            self.assertIn("needs_ocr", result.warnings)
            self.assertTrue({"ocr_unavailable", "ocr_attempted"} & set(result.warnings))
            self.assertTrue(result.is_scanned_pdf)
            self.assertIn(result.extraction_method, {"pdf_ocr", "ocr_unavailable"})
        finally:
            os.unlink(path)

    def test_html_extracts_text_and_attachments(self):
        with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tmp:
            path = Path(tmp.name)
        try:
            path.write_text("<html><body><h1>银登公告</h1><a href='notice.pdf'>附件PDF</a>本金 100 万元</body></html>", encoding="utf-8")
            result = extract_document_text(path, "公告.html")
            self.assertEqual(result.file_type, "html")
            self.assertIn("银登公告", result.text)
            self.assertEqual(result.attachments[0]["url"], "notice.pdf")
        finally:
            os.unlink(path)

    def test_image_file_uses_ocr_or_actionable_unavailable_status(self):
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            path = Path(tmp.name)
        try:
            path.write_bytes(b"\x89PNG\r\n\x1a\n")
            result = extract_document_text(path, "公告.png")
            self.assertEqual(result.file_type, "png")
            self.assertIn(result.extraction_method, {"image_ocr", "ocr_unavailable"})
            self.assertIn(result.ocr_status, {"success", "ocr_empty", "ocr_image_failed", "missing_python_deps", "missing_tesseract", "missing_poppler"})
        finally:
            os.unlink(path)

    def test_unsupported_type_returns_actionable_error(self):
        with tempfile.NamedTemporaryFile(suffix=".exe", delete=False) as tmp:
            path = Path(tmp.name)
        try:
            with self.assertRaises(DocumentExtractionError) as ctx:
                extract_document_text(path, "合同.exe")
            self.assertEqual(ctx.exception.code, "unsupported_legal_document_type")
        finally:
            os.unlink(path)


if __name__ == "__main__":
    unittest.main()
